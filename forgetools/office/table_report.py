"""forgetools.office.table_report - Convert CSV/XLSX tables into Markdown, HTML, PDF, or DOCX reports."""
from __future__ import annotations

import argparse
import html

from forgetools._cli import make_cli
from forgetools._result import ForgeResult, Timer
from forgetools.office._utils import read_table, resolve_path, table_to_markdown


def _html_table(rows: list[list[str]], title: str) -> str:
    body = []
    for idx, row in enumerate(rows):
        tag = "th" if idx == 0 else "td"
        cells = "".join(f"<{tag}>{html.escape(cell)}</{tag}>" for cell in row)
        body.append(f"<tr>{cells}</tr>")
    return f"<!doctype html><html><head><meta charset=\"utf-8\"><title>{html.escape(title)}</title></head><body><h1>{html.escape(title)}</h1><table>{''.join(body)}</table></body></html>\n"


def run(*, input: str, output: str, output_format: str = "markdown", title: str = "Table Report", sheet: str | None = None, max_rows: int = 500, cwd: str | None = None) -> ForgeResult:
    with Timer() as t:
        input_path = resolve_path(cwd, input)
        output_path = resolve_path(cwd, output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        rows = read_table(input_path, sheet=sheet, max_rows=max_rows)

        if output_format == "markdown":
            output_path.write_text(f"# {title}\n\n{table_to_markdown(rows)}\n", encoding="utf-8")
        elif output_format == "html":
            output_path.write_text(_html_table(rows, title), encoding="utf-8")
        elif output_format == "pdf":
            try:
                from reportlab.lib.pagesizes import letter, landscape
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
                from reportlab.lib import colors
                from reportlab.lib.styles import getSampleStyleSheet
            except Exception as exc:
                return ForgeResult.failure("office.table-report", [str(exc)], t.elapsed_ms, "Install reportlab in the office MCP environment.")
            styles = getSampleStyleSheet()
            story = [Paragraph(title, styles["Title"]), Spacer(1, 12), Table(rows)]
            story[-1].setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.25, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey)]))
            SimpleDocTemplate(str(output_path), pagesize=landscape(letter), title=title).build(story)
        elif output_format == "docx":
            try:
                from docx import Document
            except Exception as exc:
                return ForgeResult.failure("office.table-report", [str(exc)], t.elapsed_ms, "Install python-docx in the office MCP environment.")
            doc = Document()
            doc.add_heading(title, level=0)
            if rows:
                table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                for r_idx, row in enumerate(rows):
                    for c_idx, value in enumerate(row):
                        table.rows[r_idx].cells[c_idx].text = value
            doc.save(output_path)
        else:
            return ForgeResult.failure("office.table-report", [f"Unsupported output format: {output_format}"], t.elapsed_ms)

        return ForgeResult.success("office.table-report", {"input": str(input_path), "output": str(output_path), "rows": len(rows), "format": output_format}, t.elapsed_ms)


def _args(p: argparse.ArgumentParser) -> None:
    p.add_argument("--input", required=True, help="CSV/TSV/XLSX input")
    p.add_argument("--output", required=True, help="Report output path")
    p.add_argument("--output-format", default="markdown", choices=["markdown", "html", "pdf", "docx"], help="Report format")
    p.add_argument("--title", default="Table Report", help="Report title")
    p.add_argument("--sheet", default=None, help="Worksheet name for XLSX inputs")
    p.add_argument("--max-rows", default=500, type=int, help="Maximum rows to read")


if __name__ == "__main__":
    make_cli("office.table-report", "Convert CSV/XLSX tables into Markdown, HTML, PDF, or DOCX reports", run, _args)
