"""forgetools.office.docx_create - Create a DOCX from Markdown, HTML, or plain text."""
from __future__ import annotations

import argparse

from forgetools._cli import make_cli
from forgetools._result import ForgeResult, Timer
from forgetools.office._utils import plain_blocks, read_text, resolve_path


def run(*, input: str, output: str, source_format: str = "markdown", title: str = "Document", cwd: str | None = None) -> ForgeResult:
    with Timer() as t:
        try:
            from docx import Document
        except Exception as exc:
            return ForgeResult.failure("office.docx-create", [str(exc)], t.elapsed_ms, "Install python-docx in the office MCP environment.")

        input_path = resolve_path(cwd, input)
        output_path = resolve_path(cwd, output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.core_properties.title = title
        doc.add_heading(title, level=0)
        for kind, value in plain_blocks(read_text(input_path), source_format):
            if kind.startswith("h"):
                doc.add_heading(value, level=int(kind[1]))
            elif kind == "bullet":
                doc.add_paragraph(value, style="List Bullet")
            elif kind == "code":
                doc.add_paragraph(value, style="No Spacing")
            else:
                doc.add_paragraph(value)
        doc.save(output_path)
        return ForgeResult.success("office.docx-create", {"input": str(input_path), "output": str(output_path), "title": title}, t.elapsed_ms)


def _args(p: argparse.ArgumentParser) -> None:
    p.add_argument("--input", required=True, help="Markdown, HTML, or text input")
    p.add_argument("--output", required=True, help="DOCX output path")
    p.add_argument("--source-format", default="markdown", choices=["markdown", "html", "text"], help="Input format")
    p.add_argument("--title", default="Document", help="Document title")


if __name__ == "__main__":
    make_cli("office.docx-create", "Create a DOCX from Markdown, HTML, or plain text", run, _args)
